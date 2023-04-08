from docx import Document
import logging
import re

from .default_parser import DefaultParser
from .models import Journal, Day, Event, MCHSDep, TouristGroups, FireDep, PoliceDep, Weather
from .v2_text_processor import V2TextProcessor

logging.basicConfig(
    level=logging.DEBUG,
    format='%(levelname)s :: %(asctime)s :: %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger('v2_parser')


class V2Parser(DefaultParser):
    VERSION = 'v2'
    DEFAULT_NAME = 'default_v2'
    FILE_FORMAT = 'docx'

    EVENTS_TABLE_ID = 1
    FIRE_DEP_TABLE_ID = 2
    POLICE_DEP_TABLE_ID = 3
    MCHS_DEP_TABLE_ID = 4
    WEATHER_TABLE_ID = 5

    tables_mapping = {
        0: None,
        EVENTS_TABLE_ID: 'events',
        FIRE_DEP_TABLE_ID: 'fire',
        POLICE_DEP_TABLE_ID: 'police',
        MCHS_DEP_TABLE_ID: 'mchs',
        WEATHER_TABLE_ID: 'temp',
    }

    text_processor: V2TextProcessor

    def __init__(self, files_dir: str, file_path: str = ''):
        super(V2Parser, self).__init__(files_dir, file_path)
        self.text_processor = V2TextProcessor()

    def read_file(self, file_path: str) -> Journal:
        logger.info(f'Start parsing file: {file_path}')

        journal = Journal(
            name=file_path,
            days=[],
            version=self.VERSION
        )
        day = Day(date='', events=[], weather=[])

        try:
            word_doc = Document(file_path)
        except Exception:
            logger.error(f'File {file_path} not readable')
            return journal

        self._parse_notes(day, journal, word_doc)

        curr_table = 0
        header_row = 0

        for table in word_doc.tables:
            for i, row in enumerate(table.rows):
                # print(f'----- ROW {i} --------')

                text = tuple(cell.text for cell in row.cells)

                # detect table in file
                if text[0] == '№ п/п' and curr_table != self.EVENTS_TABLE_ID:
                    curr_table = self.EVENTS_TABLE_ID
                    header_row = i
                elif 'Пожарная' in text[0] and curr_table != self.FIRE_DEP_TABLE_ID:
                    curr_table = self.FIRE_DEP_TABLE_ID
                    header_row = i
                elif 'Полиция' in text[0] and curr_table != self.POLICE_DEP_TABLE_ID:
                    curr_table = self.POLICE_DEP_TABLE_ID
                    header_row = i
                elif 'МЧС' in text[0] and curr_table != self.MCHS_DEP_TABLE_ID:
                    curr_table = self.MCHS_DEP_TABLE_ID
                    header_row = i
                elif 'температура' in text[2] and curr_table != self.WEATHER_TABLE_ID:
                    curr_table = self.WEATHER_TABLE_ID
                    header_row = i

                # print(f'header_row={header_row} || text={text}')


                try:
                    # resolve event from table
                    if i > header_row and curr_table == self.EVENTS_TABLE_ID:
                        datetime = []
                        for datetime_part in text[1].split('\n'):
                            if datetime_part:
                                datetime.append(datetime_part.strip())

                        event_text = text[4].strip()
                        if self.need_to_skip(event_text):
                            continue

                        day.events.append(Event(
                            date=day.date,
                            text=event_text,
                            address='',
                            datetime=' '.join(datetime),
                            organization=text[2].strip(),
                            injured_fio=text[3].strip(),
                        ))

                    # resolve FireDep data from table
                    if i == header_row + 1 and curr_table == self.FIRE_DEP_TABLE_ID:
                        day.fire_dep = FireDep(
                            fires_num=text[1],
                            dead=text[3],
                            injured=text[4],
                        )

                    # resolve PoliceDep data from table
                    if i == header_row + 1 and curr_table == self.POLICE_DEP_TABLE_ID:
                        day.police_dep = PoliceDep(
                            incidents_num=text[1].strip(),
                            notes=text[2].strip(),
                        )

                    # resolve MCHS data from table
                    if i == header_row + 2 and curr_table == self.MCHS_DEP_TABLE_ID:
                        day.mchs_dep = MCHSDep(
                            tourist_groups=TouristGroups(
                                num=text[1].strip(),
                                persons=text[2].strip(),
                            ),
                            dtp=text[3].strip(),
                            search_jobs=text[4].strip(),
                            save_jobs=text[5].strip(),
                            another=text[6].strip(),
                        )

                    # resolve Weather data from table
                    if i > header_row and curr_table == self.WEATHER_TABLE_ID:
                        day.weather.append(Weather(
                            place=text[1].strip(),
                            external_temp=text[2].strip(),
                            in_out_tube_temp=text[3].strip(),
                            in_out_tube_pressure=text[4].strip(),
                        ))
                except Exception as exc:
                    logger.error(
                        f'Failed models fill for line {text} '
                        f'curr_table: {self.tables_mapping[curr_table]}'
                    )
                    raise exc

        journal.days.append(day)

        return journal

    def need_to_skip(self, line: str):
        if not line:
            return False

        lower = line.lower()

        for word in self.text_processor.skip_words:
            if word in lower:
                return True

        return False

    def _parse_notes(self, day: Day, journal: Journal, data: Document):
        lines = [para.text for para in data.paragraphs]

        for line in lines:
            # fill person on duty
            if 'дежурный' in line:
                duty_person = line.split()[-2:]
                journal.duty_person = ' '.join(duty_person)

            # fill date
            if re.match(r'.*[0-9]*\.[0-9]*\.[0-9]*', line) and 'На' not in line:
                date_field = line.split()[2]
                day.date = self.transform_date(date_field)

        # fill notes
        journal.notes = lines[-2].strip()
